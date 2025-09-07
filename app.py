import os
import tempfile
from io import StringIO

import gradio as gr
import pandas as pd
import plotly.express as px
from openai import OpenAI

# PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

# DOCX
from docx import Document
from docx.shared import Inches

# 🔑 API Key (set dalam Secrets)
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# =========================
# AI ANALISIS
# =========================
def ask_ai(question, df_head):
    messages = [
        {"role": "system", "content": "You are an AIzz Data Analyst. Jelaskan data dalam 3–5 ayat mudah difahami, sesuai untuk pelajar."},
        {"role": "user", "content": f"Contoh data:\n{df_head}"},
        {"role": "user", "content": f"Soalan: {question}"}
    ]
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,
        temperature=0.4
    )
    return resp.choices[0].message.content.strip()

# =========================
# CHART HELPERS
# =========================
def auto_chart(df):
    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    non_numeric_cols = df.select_dtypes(exclude="number").columns.tolist()
    if not numeric_cols:
        fig = px.scatter(title="⚠️ Tiada data numerik ditemui.")
        fig.update_layout(template="plotly_dark")
        return fig

    time_like = [c for c in df.columns if c.lower() in ["tahun", "year", "date", "tarikh"]]
    if time_like:
        x = time_like[0]
        fig = px.line(df, x=x, y=numeric_cols[0], title=f"Trend {numeric_cols[0]} mengikut {x}", markers=True)
    elif non_numeric_cols:
        fig = px.bar(df, x=non_numeric_cols[0], y=numeric_cols[0], title=f"Perbandingan {numeric_cols[0]} mengikut {non_numeric_cols[0]}", text=numeric_cols[0])
    else:
        fig = px.scatter(df, x=df.columns[0], y=numeric_cols[0], title=f"Hubungan {df.columns[0]} vs {numeric_cols[0]}")
    fig.update_layout(template="plotly_dark")
    return fig

def multi_graphs(df):
    figs = []
    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    if not numeric_cols:
        return []
    x = df.columns[0]
    y = numeric_cols[0]
    f1 = px.line(df, x=x, y=y, title=f"📈 Line Chart {y} vs {x}", markers=True)
    f2 = px.bar(df, x=x, y=y, title=f"📊 Bar Chart {y} vs {x}", text=y)
    f3 = px.scatter(df, x=x, y=y, title=f"🔍 Scatter Plot {y} vs {x}")
    for f in (f1, f2, f3): f.update_layout(template="plotly_dark")
    return [f1, f2, f3]

def map_chart(df):
    geo_cols = [c for c in df.columns if c.lower() in ["negeri","state","daerah","region","country","negara"]]
    num_cols = df.select_dtypes(include="number").columns.tolist()
    if not geo_cols or not num_cols:
        return None
    geo = geo_cols[0]
    val = num_cols[0]
    if geo.lower() in ["country","negara","state"]:
        fig = px.choropleth(df, locations=geo, locationmode="country names", color=val, color_continuous_scale="Viridis", title=f"🌍 Peta Dunia: {val}")
    else:
        fig = px.choropleth(
            df,
            geojson="https://raw.githubusercontent.com/sabapathy12/geojson-malaysia/master/malaysia-states.geojson",
            featureidkey="properties.name",
            locations=geo,
            color=val,
            color_continuous_scale="Blues",
            title=f"🗺️ Peta Malaysia: {val}"
        )
        fig.update_geos(fitbounds="locations", visible=False)
    fig.update_layout(template="plotly_dark")
    return fig

# =========================
# EXPORT
# =========================
def save_images(figs):
    import plotly.io as pio
    paths = []
    for i, f in enumerate(figs,1):
        path = os.path.join(tempfile.gettempdir(), f"chart_{i}.png")
        f.write_image(path, scale=2)
        paths.append(path)
    return paths

def export_pdf(title, text, images):
    path = os.path.join(tempfile.gettempdir(), "laporan.pdf")
    doc = SimpleDocTemplate(path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1,12), Paragraph("Analisis Ringkas", styles["Heading2"])]
    for p in text.split("\n"): story.append(Paragraph(p, styles["Normal"]))
    story.append(Spacer(1,12))
    for i,p in enumerate(images,1):
        story.append(Paragraph(f"Graf {i}", styles["Heading3"]))
        story.append(RLImage(p, width=400, height=250))
        story.append(Spacer(1,12))
    doc.build(story)
    return path

def export_docx(title, text, images):
    path = os.path.join(tempfile.gettempdir(), "laporan.docx")
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_heading("Analisis Ringkas", level=2)
    for p in text.split("\n"): doc.add_paragraph(p)
    for i,p in enumerate(images,1):
        doc.add_heading(f"Graf {i}", level=3)
        doc.add_picture(p, width=Inches(6))
    doc.save(path)
    return path

# =========================
# MAIN PIPELINE
# =========================
def load_df(file, text_data, url):
    if file is not None: return pd.read_csv(file.name)
    if text_data and text_data.strip(): return pd.read_csv(StringIO(text_data))
    if url and url.strip(): return pd.read_csv(url)
    return None

def pipeline(file, text_data, url, query, title):
    try:
        df = load_df(file, text_data, url)
        if df is None: return None,None,None,None,None,"⚠️ Tiada data",None,None
    except Exception as e:
        return None,None,None,None,None,f"⚠️ Ralat baca data: {e}",None,None

    try:
        explanation = ask_ai(query or "Terangkan trend utama", df.head().to_string())
    except Exception as e:
        explanation = f"⚠️ Gagal AI: {e}"

    main = auto_chart(df)
    figs = multi_graphs(df)
    mapfig = map_chart(df)
    all_figs = [main]+figs+([mapfig] if mapfig else [])

    try:
        imgs = save_images(all_figs)
        pdf = export_pdf(title or "📊 Laporan AIzz Data Analyst", explanation, imgs)
        docx = export_docx(title or "📊 Laporan AIzz Data Analyst", explanation, imgs)
    except Exception as e:
        pdf,docx=None,None
        explanation += f"\n⚠️ Gagal export: {e}"

    return main, figs[0], figs[1], figs[2], mapfig, explanation, pdf, docx

# =========================
# UI
# =========================
with gr.Blocks(css="body {background:#0f0f10; color:#fff; font-family:Inter,sans-serif;}") as demo:
    gr.Markdown("<h1 style='text-align:center'>AIzz Data Analyst (Full: Charts + Map + PDF/DOCX)</h1>")
    title = gr.Textbox(label="Tajuk Laporan", value="Laporan AIzz Data Analyst oleh Izz")
    with gr.Tab("Upload CSV"): file = gr.File(file_types=[".csv"])
    with gr.Tab("Paste Text"): text_data = gr.Textbox(lines=8, placeholder="Paste CSV text")
    with gr.Tab("URL CSV"): url = gr.Textbox(placeholder="Link CSV")
    query = gr.Textbox(placeholder="Contoh: Terangkan trend utama...", show_label=False)
    btn = gr.Button("Analisis & Laporan", variant="primary")

    main_plot = gr.Plot()
    line_plot = gr.Plot()
    bar_plot = gr.Plot()
    scatter_plot = gr.Plot()
    map_plot = gr.Plot()
    md = gr.Markdown()
    pdf_out = gr.File()
    docx_out = gr.File()

    btn.click(pipeline, [file,text_data,url,query,title], [main_plot,line_plot,bar_plot,scatter_plot,map_plot,md,pdf_out,docx_out])

demo.launch()
